import streamlit as st
import sys
import os
import docx
import faiss
from sentence_transformers import SentenceTransformer, util
import fitz
from langchain_community.llms import Ollama
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
import torch
from transformers import AutoModelForSeq2SeqLM, AutoTokenizer

class TextExtractor:
    def __init__(self, directory):
        self.directory = directory
        self.extracted_texts = {}

    def extract_text_from_docx(self, file_path):
        try:
            doc = docx.Document(file_path)
        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            return ""

        full_text = []
        capture = False
        purpose_section = False
        procedure_section = False

        for para in doc.paragraphs:
            text = para.text.strip()
            if text.startswith("PURPOSE"):
                capture = True
                purpose_section = True
                procedure_section = False
            if text.startswith("PROCEDURE"):
                capture = True
                procedure_section = True
                purpose_section = False
            if text.startswith("REFERENCES") or text.startswith("POLICY OWNER"):
                capture = False
            if capture:
                if purpose_section:
                    full_text.append(f"PURPOSE: {text}")
                    purpose_section = False  # To avoid repeating the section name
                elif procedure_section:
                    full_text.append(f"PROCEDURE: {text}")
                    procedure_section = False  # To avoid repeating the section name
                else:
                    full_text.append(text)

        return "\n".join(full_text)

    def extract_text_from_pdf(self, file_path):
        try:
            doc = fitz.open(file_path)
        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            return ""

        full_text = []
        for page in doc:
            full_text.append(page.get_text())
        return "\n".join(full_text)

    def extract_all_texts(self):
        for filename in os.listdir(self.directory):
            file_path = os.path.join(self.directory, filename)
            if filename.endswith(".docx"):
                self.extracted_texts[filename] = self.extract_text_from_docx(file_path)
            elif filename.endswith(".pdf"):
                self.extracted_texts[filename] = self.extract_text_from_pdf(file_path)
        return self.extracted_texts


class FAISS:
    def __init__(self, extracted_texts, model_name='paraphrase-MiniLM-L6-v2'):
        self.model = SentenceTransformer(model_name)
        self.texts = list(extracted_texts.values())
        self.embeddings = self.model.encode(self.texts, convert_to_tensor=True)
        self.embeddings_np = self.embeddings.cpu().numpy()

        dimension = self.embeddings_np.shape[1]
        self.index = faiss.IndexFlatL2(dimension)
        self.index.add(self.embeddings_np)

    def search(self, user_prompt, similarity_threshold=0.7):
        query_embedding = self.model.encode([user_prompt], convert_to_tensor=True).cpu().numpy()
        D, I = self.index.search(query_embedding, len(self.texts))

        # Convert distances to similarities
        similarities = 1 - D[0] / 2

        # Filter based on similarity threshold
        relevant_indices = [index for index, similarity in enumerate(similarities) if
                            similarity >= similarity_threshold]
        relevant_texts = [self.texts[index] for index in relevant_indices]

        return relevant_texts, similarities[relevant_indices]



# Paths to your logo and icon
logo_path = "childrens-hospital-la-logo.png"
icon_path = "childrens-hospital-la-icon.jpg"

import streamlit as st
import sys
import os
import time
import docx
import faiss
from sentence_transformers import SentenceTransformer, util
import fitz
from langchain_community.llms import Ollama
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
import torch
from transformers import AutoModelForSeq2SeqLM, AutoTokenizer

# Logging function to track time
def log_time(start_time, operation):
    elapsed_time = time.time() - start_time
    print(f"Time taken for {operation}: {elapsed_time:.2f} seconds")

class TextExtractor:
    def __init__(self, directory):
        self.directory = directory
        self.extracted_texts = {}

    def extract_text_from_docx(self, file_path):
        try:
            doc = docx.Document(file_path)
        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            return ""

        full_text = []
        capture = False
        purpose_section = False
        procedure_section = False

        for para in doc.paragraphs:
            text = para.text.strip()
            if text.startswith("PURPOSE"):
                capture = True
                purpose_section = True
                procedure_section = False
            if text.startswith("PROCEDURE"):
                capture = True
                procedure_section = True
                purpose_section = False
            if text.startswith("REFERENCES") or text.startswith("POLICY OWNER"):
                capture = False
            if capture:
                if purpose_section:
                    full_text.append(f"PURPOSE: {text}")
                    purpose_section = False  # To avoid repeating the section name
                elif procedure_section:
                    full_text.append(f"PROCEDURE: {text}")
                    procedure_section = False  # To avoid repeating the section name
                else:
                    full_text.append(text)

        return "\n".join(full_text)

    def extract_text_from_pdf(self, file_path):
        try:
            doc = fitz.open(file_path)
        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            return ""

        full_text = []
        for page in doc:
            full_text.append(page.get_text())
        return "\n".join(full_text)

    def extract_all_texts(self):
        start_time = time.time()
        for filename in os.listdir(self.directory):
            file_path = os.path.join(self.directory, filename)
            if filename.endswith(".docx"):
                self.extracted_texts[filename] = self.extract_text_from_docx(file_path)
            elif filename.endswith(".pdf"):
                self.extracted_texts[filename] = self.extract_text_from_pdf(file_path)
        log_time(start_time, "Text Extraction")
        return self.extracted_texts

class FAISS:
    def __init__(self, extracted_texts, model_name='paraphrase-MiniLM-L6-v2'):
        self.model = SentenceTransformer(model_name)
        self.texts = list(extracted_texts.values())
        self.embeddings = self.model.encode(self.texts, convert_to_tensor=True)
        self.embeddings_np = self.embeddings.cpu().numpy()

        dimension = self.embeddings_np.shape[1]
        self.index = faiss.IndexFlatL2(dimension)
        self.index.add(self.embeddings_np)

    def search(self, user_prompt, similarity_threshold=0.7):
        start_time = time.time()
        query_embedding = self.model.encode([user_prompt], convert_to_tensor=True).cpu().numpy()
        D, I = self.index.search(query_embedding, len(self.texts))

        # Convert distances to similarities
        similarities = 1 - D[0] / 2

        # Filter based on similarity threshold
        relevant_indices = [index for index, similarity in enumerate(similarities) if
                            similarity >= similarity_threshold]
        relevant_texts = [self.texts[index] for index in relevant_indices]

        log_time(start_time, "Vector Search")
        return relevant_texts, similarities[relevant_indices]

# Paths to your logo and icon
logo_path = "childrens-hospital-la-logo.png"
icon_path = "childrens-hospital-la-icon.jpg"

# Initialize the Ollama model
start_time = time.time()
ollama_llm = Ollama(model="llama3", base_url="http://localhost:11434")
log_time(start_time, "Ollama Initialization")

# Initialize TextExtractor
start_time = time.time()
extractor = TextExtractor(directory="/Users/andrewmorris/PycharmProjects/CHLA-LLM-Capstone-Project/sample")
extracted_texts = extractor.extract_all_texts()
log_time(start_time, "Text Extraction Initialization")

# Initialize ChromaVectorSearch
start_time = time.time()
vector_search = FAISS(extracted_texts)
log_time(start_time, "Vector Search Initialization")

# Define the updated prompt template
prompt_template = PromptTemplate.from_template("""
Context: {context}

User Question: {input_text}

Please provide a detailed and natural-sounding answer based on the context above. Maintain all medical terminology and ensure the response is clear and concise.

Answer:
""")

# Create the LLMChain
chain = LLMChain(llm=ollama_llm, prompt=prompt_template)

def generate_response(user_prompt):
    try:
        # Retrieve relevant context
        relevant_texts, _ = vector_search.search(user_prompt, similarity_threshold=st.session_state.slider)
        combined_context = "\n".join(relevant_texts)

        # Generate the combined prompt
        combined_prompt = prompt_template.format(context=combined_context, input_text=user_prompt)

        # Run the chain with the combined prompt
        start_time = time.time()
        response = chain.run({"context": combined_context, "input_text": user_prompt})
        log_time(start_time, "LLMChain Execution")
        return response
    except Exception as e:
        print(f"Error in generate_response: {e}")
        return "Apologies for the inconvenience. System still in progress. Come back soon!"

# Boot chat interface
def boot():
    # Display logo
    st.image(logo_path, width=300)

    # Title and description
    st.title("CHLA Chatbot Prototype")
    st.write("Welcome to the Children's Hospital Los Angeles Chatbot Prototype. Ask a question regarding CHLA policy!")

    # Initialize session
    if "messages" not in st.session_state:
        st.session_state.messages = []

    # Similarity threshold for testing
    st.session_state.slider = st.slider("Similarity Threshold (For Testing Purposes)", 0.0, 1.0, 0.7)

    # Display the conversation history
    for message in st.session_state.messages:
        st.chat_message("human").write(message[0])
        st.chat_message("ai").write(message[1])

    # User input for query
    if query := st.chat_input("Type your message..."):
        st.session_state.messages.append([query, ""])  # Placeholder for bot response
        st.chat_message("human").write(query)
        bot_response = generate_response(query)
        st.session_state.messages[-1][1] = bot_response
        st.chat_message("ai").write(bot_response)

if __name__ == "__main__":
    boot()

# Display the icon at the bottom
st.image(icon_path, width=100)

